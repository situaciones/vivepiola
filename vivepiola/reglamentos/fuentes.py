"""
Extraccion de texto de las fuentes normativas: PDF, Word y paginas web.

Cada formato trae su propia trampa. Un PDF escaneado no tiene texto. Un Word
guarda el contenido en parrafos y tablas separados. Una pagina web viene con
menus, banners y pies de pagina que no son la norma. Aqui se resuelve eso una
vez, para que el resto del sistema reciba texto plano y ya.
"""

import io
import re

import requests

# Un documento normativo real tiene miles de caracteres. Por debajo de esto no
# hay nada que indexar: casi siempre es un escaneo sin OCR o un archivo que no
# era el que se queria subir.
MINIMO_CARACTERES = 200

# Cabeceras de navegador: varios sitios oficiales rechazan peticiones sin ellas.
CABECERAS = {
    'User-Agent': (
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
        '(KHTML, like Gecko) Chrome/120.0 Safari/537.36'
    ),
}


class FuenteIlegible(Exception):
    """La fuente no entrega texto utilizable."""


def _de_pdf(archivo):
    import pdfplumber

    with pdfplumber.open(archivo) as pdf:
        return '\n'.join(pagina.extract_text() or '' for pagina in pdf.pages)


def _de_docx(archivo):
    from docx import Document

    documento = Document(archivo)
    partes = [p.text for p in documento.paragraphs]
    # Las tablas suelen traer los cuadros de montos y plazos, que es justo lo
    # que importa en una norma. Ignorarlas seria perder lo mas concreto.
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(' | '.join(celdas))
    return '\n'.join(partes)


def _de_url(url):
    respuesta = requests.get(url, headers=CABECERAS, timeout=30)
    respuesta.raise_for_status()

    tipo = respuesta.headers.get('Content-Type', '')
    if 'pdf' in tipo.lower() or url.lower().endswith('.pdf'):
        return _de_pdf(io.BytesIO(respuesta.content))

    from lxml import html as lxml_html

    arbol = lxml_html.fromstring(respuesta.content)
    # Fuera lo que no es la norma: menus, scripts, banners y pies.
    for etiqueta in arbol.xpath('//script | //style | //nav | //header | //footer | //aside'):
        etiqueta.getparent().remove(etiqueta)

    texto = arbol.text_content()
    # El HTML deja cascadas de saltos y espacios que inflan el texto sin aportar.
    texto = re.sub(r'[ \t]+', ' ', texto)
    return re.sub(r'\n\s*\n+', '\n\n', texto)


def extraer_texto(archivo=None, url='', nombre=''):
    """
    Devuelve el texto plano de una fuente, o levanta FuenteIlegible con un
    mensaje accionable.

    Nunca devuelve vacio en silencio: una fuente sin texto no se puede indexar
    y quien la subio tiene que enterarse en el momento, no cuando alguien note
    que la IA nunca la cita.
    """
    nombre = (nombre or getattr(archivo, 'name', '') or url).lower()

    try:
        if url and archivo is None:
            texto = _de_url(url)
        elif nombre.endswith('.pdf'):
            texto = _de_pdf(archivo)
        elif nombre.endswith(('.docx', '.doc')):
            texto = _de_docx(archivo)
        elif nombre.endswith(('.txt', '.md')):
            datos = archivo.read()
            texto = datos.decode('utf-8', errors='replace') if isinstance(datos, bytes) else datos
        else:
            raise FuenteIlegible(
                f'Formato no reconocido ({nombre or "sin nombre"}). '
                'Se admiten PDF, Word (.docx), texto plano o una URL.'
            )
    except FuenteIlegible:
        raise
    except Exception as exc:
        raise FuenteIlegible(f'No se pudo leer la fuente: {exc}') from exc

    texto = (texto or '').strip()
    if len(texto) < MINIMO_CARACTERES:
        raise FuenteIlegible(
            'La fuente no contiene texto legible. Si es un PDF escaneado como imagen, '
            'pasalo por un OCR antes de subirlo.'
        )
    return texto


# ------------------------------------------------------------------- troceado

# Los textos legales chilenos numeran sus articulos de varias formas. Cortar
# por articulo y no por largo fijo importa: el articulo es la unidad que se
# cita, y partirlo por la mitad produce fragmentos que no se pueden invocar.
PATRON_ARTICULO = re.compile(
    r'^\s*(Art[íi]culo|ART[ÍI]CULO|Art\.)\s*(\d+[°ºa-zA-Z\-]*)',
    re.MULTILINE,
)

# Tope por fragmento. Un articulo larguisimo se parte igual, porque si no
# desplaza a todos los demas al recuperar.
MAX_CARACTERES_FRAGMENTO = 3000
# Deliberadamente bajo: "Articulo 5. Prohibese fumar en espacios comunes" son
# 46 caracteres y es una norma perfectamente citable. Un umbral alto descartaba
# justo los articulos mas tajantes, que suelen ser los mas cortos.
MIN_CARACTERES_FRAGMENTO = 25


def _partir_largo(texto, referencia):
    """Parte un bloque demasiado largo respetando los parrafos."""
    parrafos = texto.split('\n\n')

    # Un parrafo que ya excede el maximo no se puede acomodar: hay que cortarlo
    # aunque quede feo. Sin esto, un articulo largo escrito de corrido nunca se
    # partia y terminaba desplazando a todos los demas al recuperar.
    sueltos = []
    for parrafo in parrafos:
        while len(parrafo) > MAX_CARACTERES_FRAGMENTO:
            corte = parrafo.rfind(' ', 0, MAX_CARACTERES_FRAGMENTO) or MAX_CARACTERES_FRAGMENTO
            sueltos.append(parrafo[:corte].strip())
            parrafo = parrafo[corte:].strip()
        if parrafo:
            sueltos.append(parrafo)

    piezas, actual = [], ''
    for parrafo in sueltos:
        if len(actual) + len(parrafo) + 2 > MAX_CARACTERES_FRAGMENTO and actual:
            piezas.append(actual.strip())
            actual = parrafo
        else:
            actual = f'{actual}\n\n{parrafo}' if actual else parrafo
    if actual.strip():
        piezas.append(actual.strip())

    if len(piezas) == 1:
        return [(referencia, piezas[0])]
    return [(f'{referencia} (parte {i})', p) for i, p in enumerate(piezas, start=1)]


def trocear(texto):
    """
    Corta el texto en fragmentos citables.

    Devuelve [(referencia, texto)]. La referencia es lo que despues aparece en
    el fundamento de una sancion, asi que se prefiere el numero de articulo por
    sobre cualquier indice interno: "Art. 12" le dice algo a una persona,
    "fragmento 47" no le dice nada a nadie.
    """
    marcas = list(PATRON_ARTICULO.finditer(texto))

    if not marcas:
        # Sin articulado (una circular, un folleto): se trocea por bloques.
        return _partir_largo(texto, 'Seccion')

    fragmentos = []
    preambulo = texto[:marcas[0].start()].strip()
    if len(preambulo) >= MIN_CARACTERES_FRAGMENTO:
        fragmentos.extend(_partir_largo(preambulo, 'Preambulo'))

    for i, marca in enumerate(marcas):
        fin = marcas[i + 1].start() if i + 1 < len(marcas) else len(texto)
        cuerpo = texto[marca.start():fin].strip()
        if len(cuerpo) < MIN_CARACTERES_FRAGMENTO:
            continue
        referencia = f'Art. {marca.group(2)}'
        if len(cuerpo) > MAX_CARACTERES_FRAGMENTO:
            fragmentos.extend(_partir_largo(cuerpo, referencia))
        else:
            fragmentos.append((referencia, cuerpo))

    return fragmentos
